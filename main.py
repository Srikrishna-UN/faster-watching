import os
from dotenv import load_dotenv
import googleapiclient.discovery
import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document

load_dotenv()

PLAYLIST_ID = os.getenv('PLAYLIST_ID')
YOUTUBE_API_KEY = os.getenv('YOUTUBE_API_KEY')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
OUTPUT_PATH = os.getenv('OUTPUT_PATH')
OUTPUT_FORMAT = os.getenv('OUTPUT_FORMAT')

SUMMARY_FILE_NAME = 'Faster watching'
PROCESSED_FILE = 'processed_videos.txt'

youtube = googleapiclient.discovery.build('youtube', 'v3', developerKey=YOUTUBE_API_KEY)
genai.configure(api_key=GEMINI_API_KEY)


# ------------------- Checkpoint Handling -------------------

def load_processed_ids():
    try:
        with open(PROCESSED_FILE, 'r', encoding='utf-8') as f:
            return set(line.strip() for line in f if line.strip())
    except FileNotFoundError:
        return set()

def save_processed_ids(processed_ids):
    with open(PROCESSED_FILE, 'w', encoding='utf-8') as f:
        for vid in sorted(processed_ids):
            f.write(vid + '\n')


# ------------------- Fetch Videos -------------------

def get_new_videos():
    processed_ids = load_processed_ids()
    videos, page_token = [], None

    while True:
        request = youtube.playlistItems().list(
            part='snippet,contentDetails',
            playlistId=PLAYLIST_ID,
            maxResults=50,
            pageToken=page_token
        )
        response = request.execute()

        for item in response.get('items', []):
            video_id = item['snippet']['resourceId']['videoId']
            title = item['snippet'].get('title', '')
            published_at = item['snippet'].get('publishedAt', '')
            videos.append({'id': video_id, 'title': title, 'publishedAt': published_at})

        page_token = response.get('nextPageToken')
        if not page_token:
            break

    # Sort oldest → newest
    videos.sort(key=lambda v: v['publishedAt'])

    # Filter only unprocessed
    new_videos = [v for v in videos if v['id'] not in processed_ids]
    return new_videos, processed_ids


# ------------------- Transcript + Summarize -------------------

def get_transcript(video_id):
    try:
        ytt_api = YouTubeTranscriptApi()
        transcript = ytt_api.fetch(video_id)
        raw_data = transcript.to_raw_data()
        full_text = " ".join([d['text'] for d in raw_data])
        return full_text
    except Exception as e:
        print(f"Could not get transcript for video {video_id}: {e}")
        return None

def summarize_with_gemini(transcript):
    model = genai.GenerativeModel('gemini-2.5-flash')
    prompt = f"""
You are a highly effective summarization assistant. I will provide a full video transcript. Your tasks are:

Write a concise, objective summary of the video's main ideas (no opinions, no unnecessary detail).

Provide a bulleted list of actionable takeaways — specific lessons, tips, or steps a viewer could apply right away.

Keep the summary concise and short unless the transcript is unusually long. Ensure the takeaways are clear, practical, and phrased so they can be implemented directly.


Transcript:
{transcript}
"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Error calling Gemini API: {e}")
        return None


def write_markdown_file(video_title, video_url, summary, video_id):
    filepath = os.path.join(OUTPUT_PATH, f"{SUMMARY_FILE_NAME}.md")
    header = """# Faster Watching Summaries

Legend:  
- Unread  
- ✅ Read  

---
"""
    entry = f"""
## {video_title}  
**[Watch on YouTube]({video_url})**

{summary}

<!-- video_id:{video_id} -->

---
"""
    os.makedirs(OUTPUT_PATH, exist_ok=True)

    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            old_content = f.read()

        if "Legend:" in old_content:
            split_point = old_content.find('---', old_content.find('Legend')) + 3
            new_content = old_content[:split_point] + entry + old_content[split_point:]
        else:
            new_content = header + entry + old_content
    else:
        new_content = header + entry

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(new_content)

    print(f"Prepended summary to: {filepath}")


def write_docx_file(video_title, video_url, summary, video_id):
    filepath = os.path.join(OUTPUT_PATH, f"{SUMMARY_FILE_NAME}.docx")

    if os.path.exists(filepath):
        doc = Document(filepath)
        first_paragraph = doc.paragraphs[0].text if doc.paragraphs else ""
        if "Legend:" not in first_paragraph:
            doc = Document()
            doc.add_paragraph("Faster Watching Summaries\n\nLegend:\n- Unread\n- ✅ Read\n\n")
    else:
        doc = Document()
        doc.add_paragraph("Faster Watching Summaries\n\nLegend:\n- Unread\n- ✅ Read\n\n")

    temp_doc = Document()
    temp_doc.add_paragraph(video_title)
    temp_doc.add_paragraph(f"Watch on YouTube: {video_url}")
    for line in summary.splitlines():
        if line.strip().startswith(('*', '-', '+')):
            temp_doc.add_paragraph(line.strip().lstrip('*-+').strip(), style='List Bullet')
        else:
            temp_doc.add_paragraph(line)
    run = temp_doc.add_paragraph().add_run(f"video_id:{video_id}")
    run.font.hidden = True
    temp_doc.add_page_break()

    for element in reversed(temp_doc._element.body):
        doc._element.body.insert(1, element)

    os.makedirs(OUTPUT_PATH, exist_ok=True)
    doc.save(filepath)
    print(f"Prepended summary to: {filepath}")


def write_txt_file(video_title, video_url, summary, video_id):
    filepath = os.path.join(OUTPUT_PATH, f"{SUMMARY_FILE_NAME}.txt")
    header = """Faster Watching Summaries

Legend:
- Unread
- ✅ Read

---
"""
    entry = f"""
{video_title}
Watch on YouTube: {video_url}

{summary}

[video_id:{video_id}]

---
"""
    os.makedirs(OUTPUT_PATH, exist_ok=True)

    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            old_content = f.read()

        if "Legend:" in old_content:
            split_point = old_content.find('---', old_content.find('Legend')) + 3
            new_content = old_content[:split_point] + entry + old_content[split_point:]
        else:
            new_content = header + entry + old_content
    else:
        new_content = header + entry

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(new_content)
    print(f"Prepended summary to: {filepath}")


def main():
    print("Checking for new videos...")
    new_videos, processed_ids = get_new_videos()
    if not new_videos:
        print("No new videos found.")
        return

    for video in new_videos:
        video_id, video_title = video['id'], video['title']
        video_url = f"https://www.youtube.com/watch?v={video_id}"
        print(f"Processing video: {video_title} ({video_url})")

        transcript = get_transcript(video_id)
        if not transcript:
            print("Skipping due to failed transcript retrieval.")
            continue

        summary = summarize_with_gemini(transcript)
        if not summary:
            print("Failed to generate summary.")
            continue

        try:
            if OUTPUT_FORMAT == 'md':
                write_markdown_file(video_title, video_url, summary, video_id)
            elif OUTPUT_FORMAT == 'docx':
                write_docx_file(video_title, video_url, summary, video_id)
            elif OUTPUT_FORMAT == 'txt':
                write_txt_file(video_title, video_url, summary, video_id)
            else:
                print(f"Error: Invalid OUTPUT_FORMAT in .env: {OUTPUT_FORMAT}")
                break
        except Exception as e:
            print(f"Failed to write summary for {video_id}: {e}")
            continue

        processed_ids.add(video_id)
        save_processed_ids(processed_ids)


if __name__ == "__main__":
    main()
